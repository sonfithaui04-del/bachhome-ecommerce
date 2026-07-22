
import os
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Installing python-docx...")
    os.system('pip install python-docx')
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)

    # --- TITLE PAGE ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ĐẠI HỌC CÔNG NGHIỆP HÀ NỘI\nTRƯỜNG CÔNG NGHỆ THÔNG TIN VÀ TRUYỀN THÔNG\n=====***=====\n\n\n\n')
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('BÁO CÁO\nĐỒ ÁN TỐT NGHỆP\n\n')
    run.bold = True
    run.font.size = Pt(20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('XÂY DỰNG HỆ THỐNG ĐẶT MÓN ĂN TRỰC TUYẾN SỬ DỤNG KIẾN TRÚC MICROSERVICES VỚI MÔ HÌNH API GATEWAY\n\n\n')
    run.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('GVHD: TS. Hà Mạnh Đào\nLớp: 20252IT6009001\nHọ và tên: Đỗ Thiên Nhật\nMã sinh viên: 2022603404\n\n\n\n\n')
    run.font.size = Pt(13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('HÀ NỘI - 2026')
    
    doc.add_page_break()

    # --- LỜI NÓI ĐẦU ---
    doc.add_heading('LỜI NÓI ĐẦU', level=1)
    doc.add_paragraph("Trong bối cảnh công nghệ thông tin phát triển mạnh mẽ, các hệ thống phần mềm hiện đại ngày càng đòi hỏi cao về khả năng mở rộng, tính linh hoạt và hiệu năng...")
    
    doc.add_page_break()

    # --- CHƯƠNG 1 ---
    doc.add_heading('CHƯƠNG 1. CƠ SỞ LÝ THUYẾT', level=1)
    doc.add_heading('1.1. Tổng quan kiến trúc phần mềm hướng dịch vụ', level=2)
    doc.add_paragraph("Kiến trúc hướng dịch vụ (SOA) và Microservices là những nền tảng quan trọng trong phát triển phần mềm hiện đại...")

    # --- CHƯƠNG 2 ---
    doc.add_heading('CHƯƠNG 2. PHÂN TÍCH VÀ ĐẶC TẢ YÊU CẦU BÀI TOÁN', level=1)
    doc.add_heading('2.2. Yêu cầu chức năng bổ sung', level=2)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'STT'
    hdr_cells[1].text = 'Module'
    hdr_cells[2].text = 'Chức năng'
    
    data = [
        ('7', 'Review & Rating', 'Khách hàng đánh giá (sao) và viết bình luận cho món ăn.'),
        ('8', 'Real-time Chat', 'Chat trực tiếp giữa Khách hàng và Admin qua WebSocket.'),
        ('9', 'QR Payment', 'Tích hợp SePay tự động tạo QR Code và Webhook xác nhận.')
    ]
    
    for stt, mod, func in data:
        row_cells = table.add_row().cells
        row_cells[0].text = stt
        row_cells[1].text = mod
        row_cells[2].text = func

    # --- CHƯƠNG 3 ---
    doc.add_heading('CHƯƠNG 3. THIẾT KẾ HỆ THỐNG', level=1)
    doc.add_heading('3.1.3. Domain Model bổ sung', level=2)
    doc.add_paragraph("- Review Domain: id, orderId, menuItemId, rating (1-5), comment, createdAt.")
    doc.add_paragraph("- Chat Domain: MessageContent, SenderId, ReceiverId, Timestamp.")

    # --- CHƯƠNG 4 ---
    doc.add_heading('CHƯƠNG 4. CÀI ĐẶT VÀ TRIỂN KHAI', level=1)
    doc.add_heading('4.3. Cấu hình Thanh toán SePay và WebSocket', level=2)
    doc.add_paragraph("Hệ thống tích hợp cổng SePay giúp tự động hóa thanh toán qua QR. Giao thức WebSocket được triển khai qua API Gateway để hỗ trợ tính năng Chat thời gian thực.")

    doc.add_heading('KẾT LUẬN', level=1)
    doc.add_paragraph("Đồ án đã hoàn thành xuất sắc việc xây dựng hệ thống Microservices chuẩn chỉnh, tích hợp các tính năng hiện đại nhất hiện nay.")

    # Save
    file_name = "Bao_Cao_Do_An_Tot_Nghiep_Final.docx"
    doc.save(file_name)
    print(f"Successfully exported to {file_name}")

if __name__ == "__main__":
    create_report()
